import docx
import random

# Function to generate a random addition question
def generate_addition_question():
    # Generate two random numbers between 1 and 9
    num1 = random.randint(1, 9)
    num2 = random.randint(1, 9)

    # Create the question string
    question = f"{num1} + {num2} = __"

    # Add a space at the end to match the format in the image
    question += " "

    return question

# Create a Word document
doc = docx.Document()

# Add 100 random addition questions to the document
for _ in range(100):
    question = generate_addition_question()
    doc.add_paragraph(question)
    doc.add_paragraph("\n")  # Add some space between questions

# Save the document
doc.save("random_addition_questions.docx")

print("Document saved successfully.")
